import telebot
from telebot import types
import os
import tempfile
import smtplib
from email.message import EmailMessage
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Mm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

bot = telebot.TeleBot("7955924179:AAFiaz3iF4nSfx7PNaW9jNPnasidu1zYZ1k")

user_states = {}

questions = [
    ("sender_name", "Введите ФИО отправителя:"),
    ("company_name", "Введите название компании, отправляющей заявление:"),
    ("project_name", "Введите название проекта:"),
    ("engine_number", "Введите номер двигателя:"),
    ("project_number", "Введите номер проекта:"),
    ("unit_number", "Введите номер агрегата:"),
    ("moto_hours", "Введите наработку мотто-часов:"),
    ("start_count", "Введите количество стартов:"),
    ("problem_description", "Опишите проблему:")
]

SPARE_PARTS_STAGE = 'spare_parts_stage'

SMTP_SERVER = "smtp-mail.outlook.com"  # Замените на ваш SMTP сервер
SMTP_PORT = 587  # Обычно 587 для TLS
SMTP_USER = "warranty@gte.su"  # Ваш email
SMTP_PASSWORD = "7d2758Pz7"  # Ваш пароль или app password


# ============================

@bot.message_handler(commands=['start'])
def start(message):
    user_states[message.from_user.id] = {"step": 0, "data": {}}
    # Список вопросов
    question_list = '\n'.join([f"{i + 1}. {q[1]}" for i, q in enumerate(questions)])
    markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
    btn1 = types.KeyboardButton("Начать заполнение")
    markup.add(btn1)
    full_text = (
            "Перед заполнением ознакомьтесь со списком вопросов:\n\n"
            + question_list +
            "\n\nНажмите 'Начать заполнение' для подачи заявления."
    )
    bot.send_message(message.from_user.id, full_text, reply_markup=markup)


@bot.message_handler(func=lambda message: message.text == "Начать заполнение")
def begin_form(message):
    user_states[message.from_user.id] = {"step": 0, "data": {}}
    ask_next_question(message)


def ask_next_question(message):
    user_id = message.from_user.id
    state = user_states.get(user_id)
    if state is None:
        return
    step = state["step"]
    if step < len(questions):
        key, question = questions[step]
        markup = types.ReplyKeyboardMarkup(resize_keyboard=True, one_time_keyboard=True)
        skip_btn = types.KeyboardButton("Пропустить")
        markup.add(skip_btn)
        bot.send_message(user_id, question, reply_markup=markup)
    else:
        state[SPARE_PARTS_STAGE] = {"stage": 0, "current": {}, "parts": []}
        bot.send_message(user_id, "Введите данные о запасных частях. Если не требуется — нажмите 'Пропустить'.")
        ask_spare_part(message)


def ask_spare_part(message):
    user_id = message.from_user.id
    state = user_states.get(user_id)
    sp = state.get(SPARE_PARTS_STAGE)
    if not sp:
        return
    stage = sp["stage"]
    markup = types.ReplyKeyboardMarkup(resize_keyboard=True, one_time_keyboard=True)
    skip_btn = types.KeyboardButton("Пропустить")
    markup.add(skip_btn)
    if stage == 0:
        bot.send_message(user_id, "Каталожный номер:", reply_markup=markup)
    elif stage == 1:
        bot.send_message(user_id, "Наименование:", reply_markup=markup)
    elif stage == 2:
        bot.send_message(user_id, "Количество:", reply_markup=markup)
    elif stage == 3:
        markup = types.ReplyKeyboardMarkup(resize_keyboard=True, one_time_keyboard=True)
        yes_btn = types.KeyboardButton("Добавить ещё")
        no_btn = types.KeyboardButton("Закончить ввод")
        markup.add(yes_btn, no_btn)
        part = sp["current"]
        bot.send_message(user_id,
                         f"Добавить ещё одну позицию?\nТекущая: {part['catalog']} | {part['name']} | {part['qty']}",
                         reply_markup=markup)


@bot.message_handler(
    func=lambda message: (
            user_states.get(message.from_user.id, {}).get("step") is not None
            and user_states.get(message.from_user.id, {}).get("step") < len(questions)
            and SPARE_PARTS_STAGE not in user_states.get(message.from_user.id, {})
            and not user_states.get(message.from_user.id, {}).get("attaching_files")
    )
)
def handle_answers(message):
    user_id = message.from_user.id
    state = user_states.get(user_id)
    if state is None:
        return
    step = state["step"]
    key, _ = questions[step]
    answer = message.text.strip()
    if answer.lower() == "пропустить" or answer == "-":
        answer = ""
    state["data"][key] = answer
    state["step"] += 1
    ask_next_question(message)


@bot.message_handler(func=lambda message: SPARE_PARTS_STAGE in user_states.get(message.from_user.id, {}))
def handle_spare_parts(message):
    user_id = message.from_user.id
    state = user_states.get(user_id)
    sp = state.get(SPARE_PARTS_STAGE)
    if not sp:
        return
    text = message.text.strip()
    if text.lower() == "пропустить" or text == "-":
        state["data"]["spare_parts"] = []
        state.pop(SPARE_PARTS_STAGE, None)
        ask_next_question_after_spare_parts(message)
        return
    if sp["stage"] == 0:
        sp["current"] = {"catalog": text}
        sp["stage"] = 1
        ask_spare_part(message)
    elif sp["stage"] == 1:
        sp["current"]["name"] = text
        sp["stage"] = 2
        ask_spare_part(message)
    elif sp["stage"] == 2:
        sp["current"]["qty"] = text
        sp["stage"] = 3
        ask_spare_part(message)
    elif sp["stage"] == 3:
        if text.lower() == "добавить ещё":
            sp["parts"].append(sp["current"])
            sp["current"] = {}
            sp["stage"] = 0
            ask_spare_part(message)
        elif text.lower() == "закончить ввод":
            sp["parts"].append(sp["current"])
            state["data"]["spare_parts"] = sp["parts"]
            state.pop(SPARE_PARTS_STAGE, None)
            ask_next_question_after_spare_parts(message)
        else:
            ask_spare_part(message)


def ask_next_question_after_spare_parts(message):
    user_id = message.from_user.id
    state = user_states.get(user_id)
    markup = types.ReplyKeyboardMarkup(resize_keyboard=True, one_time_keyboard=True)
    skip_btn = types.KeyboardButton("Пропустить")
    finish_btn = types.KeyboardButton("Завершить прикрепление")
    # Сначала 'Пропустить', потом 'Завершить прикрепление' (справа)
    markup.add(skip_btn, finish_btn)
    bot.send_message(user_id,
                     "Прикрепите необходимые файлы (фото или документы). Когда закончите — нажмите 'Завершить прикрепление' или 'Пропустить'.",
                     reply_markup=markup)
    state["attaching_files"] = True
    state["files"] = []


@bot.message_handler(content_types=['document', 'photo'])
def handle_files(message):
    user_id = message.from_user.id
    state = user_states.get(user_id)
    if not state or not state.get("attaching_files"):
        return
    if message.content_type == 'document':
        file_id = message.document.file_id
        file_name = message.document.file_name
        state["files"].append({"type": "document", "file_id": file_id, "file_name": file_name})
        bot.reply_to(message, f"Документ '{file_name}' добавлен.")
    elif message.content_type == 'photo':
        file_id = message.photo[-1].file_id
        # Скачиваем фото во временный файл
        file_info = bot.get_file(file_id)
        downloaded_file = bot.download_file(file_info.file_path)
        temp_dir = tempfile.gettempdir()
        photo_path = os.path.join(temp_dir, f"{file_id}.jpg")
        with open(photo_path, 'wb') as new_file:
            new_file.write(downloaded_file)
        state["files"].append({"type": "photo", "file_id": file_id, "photo_path": photo_path})
        # Показываем пользователю само фото
        with open(photo_path, 'rb') as photo_file:
            bot.send_photo(user_id, photo_file, caption="Фото добавлено.")


@bot.message_handler(func=lambda message: user_states.get(message.from_user.id, {}).get("attaching_files"))
def handle_finish_attach(message):
    user_id = message.from_user.id
    state = user_states.get(user_id)
    if not state or not state.get("attaching_files"):
        return
    if message.text.strip().lower() in ["завершить прикрепление", "пропустить"]:
        state["data"]["attached_files"] = state.get("files", [])
        state.pop("attaching_files", None)
        state.pop("files", None)

        data = state["data"]
        file_path_docx = generate_docx(data)

        with open(file_path_docx, "rb") as docx_file:
            bot.send_document(user_id, docx_file)

        # Отправляем DOCX на почту
        subject = "Новое гарантийное заявление"
        body = f"Поступило новое гарантийное заявление от {data.get('company_name', '-')}, проект: {data.get('project_name', '-')}."
        send_files_to_email([file_path_docx], subject, body, "warranty@gte.su")

        bot.send_message(user_id, "Спасибо! Ваше заявление сформировано и отправлено в виде Word (docx).")
        user_states.pop(user_id, None)
        os.remove(file_path_docx)
        # Автоматический рестарт: предлагаем начать заново
        markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
        btn1 = types.KeyboardButton("Начать заполнение")
        markup.add(btn1)
        bot.send_message(user_id, "Хотите подать новое заявление? Нажмите 'Начать заполнение'.", reply_markup=markup)


def generate_docx(data):
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn
    import datetime
    from docx.shared import Mm

    doc = Document()

    # Установка полей документа
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    def styled_run(paragraph, text, bold=False, italic=False, font_size=10):
        run = paragraph.add_run(text)
        run.font.name = 'Montserrat'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Montserrat')
        run.font.size = Pt(font_size)
        run.bold = bold
        run.italic = italic
        # Одинарный межстрочный интервал
        paragraph.paragraph_format.line_spacing = 1
        return run

    def center_paragraph(text, bold=False, font_size=10):
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        styled_run(p, text, bold=bold, font_size=font_size)
        p.paragraph_format.line_spacing = 1
        return p

    # Добавляем изображение верхнего колонтитула
    doc.add_picture('верхний колонтитул.png',
                    width=doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin)

    # Исходящий номер и дата
    today = datetime.datetime.today()
    p = doc.add_paragraph()
    styled_run(p, f'Исх. №___ от «__» _____ {today.year} г.', font_size=10)
    p.paragraph_format.line_spacing = 1

    # Заголовок
    center_paragraph("Заявление", font_size=22)
    center_paragraph("о гарантийном случае от Warranty case claim", font_size=11)

    # Таблица параметров (8 строк, 3 колонки)
    table = doc.add_table(rows=9, cols=3)
    table.style = 'Table Grid'
    col_widths = [Cm(0.8), Cm(8.0), Cm(8.0)]
    params = [
        ("1", "Название проекта\nProject name", data.get("project_name", "")),
        ("2", "Номер проекта\nDesing number", data.get("project_number", "")),
        ("3", "Номер двигателя\nEngine number", data.get("engine_number", "")),
        ("4", "Номер агрегата\nUnit number", data.get("unit_number", "")),
        ("5", "Наработка моточасов\nOph", data.get("moto_hours", "")),
        ("6", "Количество стартов\nNumber of starts", data.get("start_count", "")),
        ("7", "Описание проблемы\nProblem description", data.get("problem_description", "")),
        ("8", "Прилагаемые материалы\nAttachments", None)
    ]

    table.cell(0, 0).text = "№ п/п"
    table.cell(0, 1).text = ""
    table.cell(0, 2).text = ""

    for i, (n, label, value) in enumerate(params):
        table.cell(i + 1, 0).text = n
        cell_label = table.cell(i + 1, 1).paragraphs[0]
        if '\n' in label:
            ru, en = label.split('\n')
            styled_run(cell_label, ru + "\n", bold=True, font_size=10)
            styled_run(cell_label, en, italic=True, font_size=10)
        else:
            styled_run(cell_label, label, font_size=10)
        cell_label.paragraph_format.line_spacing = 1
        cell_value = table.cell(i + 1, 2)
        paragraph = cell_value.paragraphs[0]
        if n != "8":
            styled_run(paragraph, value, font_size=10)
        else:
            files = data.get("attached_files", [])
            if files:
                for f in files:
                    if f['type'] == 'photo' and f.get('photo_path') and os.path.exists(f['photo_path']):
                        try:
                            run = paragraph.add_run()
                            run.add_picture(f['photo_path'], width=Mm(13))
                            paragraph.add_run("\n")
                        except Exception:
                            paragraph.add_run("(ошибка вставки)\n")
                    elif f['type'] == 'document':
                        styled_run(paragraph, f"• {f.get('file_name') or f['file_id']}\n", font_size=10)
            else:
                styled_run(paragraph, "—", font_size=10)
        paragraph.paragraph_format.line_spacing = 1
        # Устанавливаем ширину столбцов
        table.columns[0].width = col_widths[0]
        table.columns[1].width = col_widths[1]
        table.columns[2].width = col_widths[2]

    # Запасные части
    p = doc.add_paragraph()
    if '\n' in "Запасные части, необходимые для устранения проблемы\nSpare parts needed":
        ru, en = "Запасные части, необходимые для устранения проблемы\nSpare parts needed".split('\n')
        styled_run(p, ru + "\n", bold=True, font_size=10)
        styled_run(p, en, italic=True, font_size=10)
    else:
        styled_run(p, "Запасные части, необходимые для устранения проблемы", bold=True, font_size=10)
        styled_run(p, "Spare parts needed", font_size=10)
    p.paragraph_format.line_spacing = 1
    parts = data.get("spare_parts", [])
    if parts:
        table2 = doc.add_table(rows=1, cols=3)
        table2.style = 'Table Grid'
        hdr_cells = table2.rows[0].cells
        for idx, hdr in enumerate(["Каталожный номер\nPartial number", "Название\nDescription", "Количество\nAmount"]):
            cell = hdr_cells[idx].paragraphs[0]
            if '\n' in hdr:
                ru, en = hdr.split('\n')
                styled_run(cell, ru + "\n", bold=True, font_size=10)
                styled_run(cell, en, italic=True, font_size=10)
            else:
                styled_run(cell, hdr, bold=True, font_size=10)
        for part in parts:
            row = table2.add_row().cells
            styled_run(row[0].paragraphs[0], part["catalog"], font_size=10)
            styled_run(row[1].paragraphs[0], part["name"], font_size=10)
            styled_run(row[2].paragraphs[0], part["qty"], font_size=10)
            for cell in row:
                for p in cell.paragraphs:
                    p.paragraph_format.line_spacing = 1
    else:
        p = doc.add_paragraph()
        styled_run(p, "—", font_size=10)
        p.paragraph_format.line_spacing = 1

    doc.add_paragraph()

    # Блок подписей — 3 столбца, без границ
    table3 = doc.add_table(rows=5, cols=3)
    table3.style = None  # Без стиля, чтобы не было границ
    col1, col2, col3 = table3.columns
    col1.width = Cm(6.5)
    col2.width = Cm(5.0)
    col3.width = Cm(6.5)

    # Первая колонка — Гринтех
    styled_run(table3.cell(0, 0).paragraphs[0], 'ООО «ГринТех Энерджи»', font_size=9)
    styled_run(table3.cell(1, 0).paragraphs[0], 'GreenTech Energy LLC', font_size=9)
    table3.cell(3, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    styled_run(table3.cell(3, 0).paragraphs[0], '_____________________/______________/', font_size=9)

    # Центр — дата
    table3.cell(0, 1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    styled_run(table3.cell(0, 1).paragraphs[0], 'Дата / Date', font_size=9)

    # Правая колонка — компания пользователя
    styled_run(table3.cell(0, 2).paragraphs[0], f'«{today.day}» {today.month} {today.year} г.', font_size=9)
    table3.cell(3, 2).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    styled_run(table3.cell(3, 2).paragraphs[0], '___________________/__________________/', font_size=9)

    for row in table3.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.line_spacing = 1

    file_path = f"zayavlenie_{data.get('engine_number', 'no_engine')}.docx"
    doc.save(file_path)
    return file_path


def send_files_to_email(file_paths, subject, body, to_email):
    from_email = SMTP_USER
    if not all([SMTP_SERVER, SMTP_PORT, SMTP_USER, SMTP_PASSWORD]):
        print('SMTP credentials are not set.')
        return
    msg = EmailMessage()
    msg['Subject'] = subject
    msg['From'] = from_email
    msg['To'] = to_email
    msg.set_content(body)
    for file_path in file_paths:
        with open(file_path, 'rb') as f:
            file_data = f.read()
            file_name = os.path.basename(file_path)
        maintype = 'application'
        subtype = 'octet-stream'
        if file_name.endswith('.pdf'):
            subtype = 'pdf'
        elif file_name.endswith('.xlsx'):
            subtype = 'vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        msg.add_attachment(file_data, maintype=maintype, subtype=subtype, filename=file_name)
    try:
        with smtplib.SMTP(SMTP_SERVER, SMTP_PORT) as server:
            server.starttls()
            server.login(SMTP_USER, SMTP_PASSWORD)
            server.send_message(msg)
    except Exception as e:
        print(f'Error sending email: {e}')


if __name__ == '__main__':
    bot.polling(none_stop=True)
